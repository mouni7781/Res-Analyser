"""
Utilities for extracting plain text from PDF and DOCX resume files.
"""
import io

from logger import get_logger

logger = get_logger("extractor")


def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Returns extracted plain text.
    """
    lower = filename.lower()
    size_kb = len(content) / 1024
    logger.debug(f"extract_text_from_file: '{filename}'  size={size_kb:.1f} KB")

    if lower.endswith(".pdf"):
        text = _extract_pdf(content, filename)
    elif lower.endswith(".docx"):
        text = _extract_docx(content, filename)
    elif lower.endswith(".doc"):
        raise ValueError(
            "Legacy .doc format is not supported. Please convert to .docx or .pdf."
        )
    else:
        logger.warning(f"Unknown extension for '{filename}' — attempting UTF-8 decode")
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception as exc:
            raise ValueError(f"Unsupported file format: {filename}") from exc

    logger.debug(f"Extraction complete for '{filename}': {len(text)} chars extracted")
    return text


def _extract_pdf(content: bytes, filename: str) -> str:
    """Extract text from a PDF using pypdf."""
    try:
        import pypdf
    except ImportError as exc:
        raise ImportError("pypdf is required: pip install pypdf") from exc

    reader = pypdf.PdfReader(io.BytesIO(content))
    total_pages = len(reader.pages)
    logger.debug(f"  PDF '{filename}': {total_pages} page(s)")

    pages = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
                logger.debug(f"    Page {page_num}/{total_pages}: {len(text)} chars")
            else:
                logger.debug(f"    Page {page_num}/{total_pages}: no text (possibly image-based)")
        except Exception as exc:
            logger.warning(f"    Page {page_num}/{total_pages} extraction error: {exc}")

    if not pages:
        logger.warning(f"  No text extracted from PDF '{filename}' — may be a scanned image")

    return "\n".join(pages)


def _extract_docx(content: bytes, filename: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    doc = docx.Document(io.BytesIO(content))
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    logger.debug(f"  DOCX '{filename}': {len(paragraphs)} non-empty paragraph(s)/cell(s)")
    return "\n".join(paragraphs)
